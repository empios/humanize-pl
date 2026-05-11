from __future__ import annotations

import csv
from dataclasses import asdict, dataclass, field
import json
from pathlib import Path
from statistics import mean
import tempfile
from typing import Any

import regex as re
import typer
from rich import print

from humanize_pl.config import Engine, Mode
from humanize_pl.core import HumanizeResult, humanize_text
from humanize_pl.io.docx_io import process_docx
from humanize_pl.reports.report import write_json_report
from humanize_pl.safety.protectors import protect_text


DEFAULT_MANIFEST = Path("docs_tests/ai_generated/manifest.json")
DEFAULT_OUTPUT = Path("docs_tests/results/latest")
DEFAULT_ENGINES = "basic,hybrid"

app = typer.Typer(add_completion=False, help="Benchmark humanize-pl on legal AI fixtures.")


@dataclass(frozen=True)
class BenchmarkDocument:
    id: str
    path: Path
    type: str = "unknown"
    focus: list[str] = field(default_factory=list)
    source_kind: str = "txt"


@dataclass
class BenchmarkRow:
    document_id: str
    document_type: str
    engine: str
    mode: str
    status: str
    source_path: str
    output_path: str | None = None
    report_path: str | None = None
    accepted_changes: int = 0
    rejected_candidates: int = 0
    skipped_sentences: int = 0
    all_candidates: int = 0
    changes_per_1000_words: float = 0.0
    average_accepted_risk: float = 0.0
    model_status: dict[str, str] = field(default_factory=dict)
    semantic_model: str | None = None
    fluency_model: str | None = None
    operation_types: dict[str, int] = field(default_factory=dict)
    gate_rejections: dict[str, int] = field(default_factory=dict)
    safety: dict[str, Any] = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)
    error: str | None = None


@app.command()
def main(
    manifest: Path = typer.Option(DEFAULT_MANIFEST, "--manifest", help="Benchmark manifest JSON"),
    output: Path = typer.Option(DEFAULT_OUTPUT, "--output", "-o", help="Output directory"),
    engines: str = typer.Option(DEFAULT_ENGINES, "--engines", help="Comma-separated engines"),
    mode: Mode = typer.Option(Mode.standard, "--mode", help="conservative, standard, strong"),
    offline_models: bool = typer.Option(True, "--offline-models/--online-models", help="Use local model cache only"),
    require_models: bool = typer.Option(False, "--require-models", help="Fail if requested models are unavailable"),
    allow_fallback: bool = typer.Option(False, "--allow-fallback", help="Allow nlp/hybrid to fall back to available layers"),
    include_docx: list[Path] = typer.Option(
        [],
        "--include-docx",
        help="Additional DOCX file to benchmark",
    ),
) -> None:
    documents = load_manifest(manifest)
    documents.extend(_docx_documents(include_docx))
    selected_engines = parse_engines(engines)
    rows = run_benchmark(
        documents,
        output_dir=output,
        engines=selected_engines,
        mode=mode,
        offline_models=offline_models,
        require_models=require_models,
        allow_fallback=allow_fallback,
    )
    write_summary_artifacts(rows, output)
    print(f"[green]Benchmark zapisany:[/green] {output}")
    print(f"Dokumenty: {len(documents)}")
    print(f"Uruchomienia: {len(rows)}")
    failed = sum(1 for row in rows if row.status != "ok")
    if failed:
        print(f"[yellow]Statusy wymagające uwagi:[/yellow] {failed}")


def parse_engines(value: str) -> list[Engine]:
    engines: list[Engine] = []
    for raw in value.split(","):
        item = raw.strip()
        if not item:
            continue
        engines.append(Engine(item))
    if not engines:
        raise ValueError("at least one engine is required")
    return engines


def load_manifest(path: str | Path) -> list[BenchmarkDocument]:
    path = Path(path)
    data = json.loads(path.read_text(encoding="utf-8"))
    documents: list[BenchmarkDocument] = []
    for item in data:
        file_name = item["file"]
        source_path = path.parent / file_name
        if not source_path.exists():
            raise FileNotFoundError(f"Manifest document not found: {source_path}")
        documents.append(
            BenchmarkDocument(
                id=item["id"],
                path=source_path,
                type=item.get("type", "unknown"),
                focus=list(item.get("focus", [])),
                source_kind="txt",
            )
        )
    return sorted(documents, key=lambda doc: doc.id)


def run_benchmark(
    documents: list[BenchmarkDocument],
    *,
    output_dir: Path,
    engines: list[Engine],
    mode: Mode,
    offline_models: bool,
    require_models: bool,
    allow_fallback: bool,
) -> list[BenchmarkRow]:
    output_dir.mkdir(parents=True, exist_ok=True)
    rows: list[BenchmarkRow] = []
    for document in sorted(documents, key=lambda doc: doc.id):
        original_text = _read_document_text(document)
        for engine in engines:
            row = _run_one(
                document,
                original_text=original_text,
                output_dir=output_dir,
                engine=engine,
                mode=mode,
                offline_models=offline_models,
                require_models=require_models or (engine != Engine.basic and not allow_fallback),
            )
            rows.append(row)
    return rows


def write_summary_artifacts(rows: list[BenchmarkRow], output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    summary = {
        "summary": _aggregate(rows),
        "rows": [asdict(row) for row in rows],
    }
    (output_dir / "summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    _write_summary_csv(rows, output_dir / "summary.csv")
    (output_dir / "review.md").write_text(render_review_markdown(rows), encoding="utf-8")


def safety_checks(original: str, rewritten: str) -> dict[str, Any]:
    original_protected = protect_text(original)
    rewritten_protected = protect_text(rewritten)
    checks = {
        "no_english_markers": not re.search(
            r"\b(?:surprisingly|actually|basically|honestly|literally|overall|moreover)\b",
            rewritten,
            re.IGNORECASE,
        ),
        "no_bad_split_phrase": "Ponadto za wynagrodzeniem" not in rewritten,
        "no_placeholder_leak": "__PROTECTED_" not in rewritten,
        "numbers_preserved": _numbers(original) == _numbers(rewritten),
        "protected_fragments_preserved": _protected_values(original_protected)
        <= _protected_values(rewritten_protected),
    }
    checks["passed"] = all(checks.values())
    return checks


def render_review_markdown(rows: list[BenchmarkRow]) -> str:
    lines: list[str] = []
    lines.append("# humanize-pl Benchmark Review")
    lines.append("")
    lines.append("## Summary")
    lines.append("")
    aggregate = _aggregate(rows)
    lines.append(f"- Runs: {aggregate['runs']}")
    lines.append(f"- OK: {aggregate['ok']}")
    lines.append(f"- Failed safety: {aggregate['failed_safety']}")
    lines.append(f"- Model unavailable: {aggregate['model_unavailable']}")
    lines.append(f"- Accepted changes: {aggregate['accepted_changes']}")
    lines.append("")
    lines.append("## Per Document")
    lines.append("")
    lines.append(
        "| Document | Engine | Status | Accepted | Rejected | Risk | Changes/1000 | Safety |"
    )
    lines.append("|---|---:|---:|---:|---:|---:|---:|---:|")
    for row in sorted(rows, key=lambda item: (item.document_id, item.engine)):
        lines.append(
            "| "
            f"{row.document_id} | {row.engine} | {row.status} | {row.accepted_changes} | "
            f"{row.rejected_candidates} | {row.average_accepted_risk:.4f} | "
            f"{row.changes_per_1000_words:.4f} | {row.safety.get('passed', False)} |"
        )
    lines.append("")
    lines.append("## Rejected Candidates")
    lines.append("")
    rejected_sections = _rejected_sections(rows)
    lines.extend(rejected_sections or ["Brak odrzuconych kandydatów."])
    lines.append("")
    lines.append("## Needs Review")
    lines.append("")
    needs_review = _needs_review_sections(rows)
    lines.extend(needs_review or ["Brak zmian oznaczonych do ręcznego przeglądu."])
    lines.append("")
    lines.append("## Recommended Next Rules")
    lines.append("")
    lines.extend(_recommended_rules(rows))
    lines.append("")
    return "\n".join(lines)


def _run_one(
    document: BenchmarkDocument,
    *,
    original_text: str,
    output_dir: Path,
    engine: Engine,
    mode: Mode,
    offline_models: bool,
    require_models: bool,
) -> BenchmarkRow:
    engine_dir = output_dir / engine.value
    engine_dir.mkdir(parents=True, exist_ok=True)
    row = BenchmarkRow(
        document_id=document.id,
        document_type=document.type,
        engine=engine.value,
        mode=mode.value,
        status="ok",
        source_path=str(document.path),
    )
    try:
        result, output_path = _process_document(
            document,
            engine_dir=engine_dir,
            engine=engine,
            mode=mode,
            offline_models=offline_models,
            require_models=require_models,
        )
    except RuntimeError as exc:
        row.status = "model_unavailable"
        row.error = str(exc)
        return row

    report_path = engine_dir / f"{document.id}.json"
    write_json_report(result, report_path)
    payload = json.loads(report_path.read_text(encoding="utf-8"))
    safety = safety_checks(original_text, result.text if document.source_kind == "txt" else _read_docx_text(output_path))
    status = "ok" if safety["passed"] else "failed_safety"
    return _row_from_payload(
        row,
        payload=payload,
        result=result,
        output_path=output_path,
        report_path=report_path,
        safety=safety,
        status=status,
    )


def _process_document(
    document: BenchmarkDocument,
    *,
    engine_dir: Path,
    engine: Engine,
    mode: Mode,
    offline_models: bool,
    require_models: bool,
) -> tuple[HumanizeResult, Path]:
    if document.source_kind == "docx":
        output_path = engine_dir / f"{document.id}.docx"
        result, _stats = process_docx(
            document.path,
            output_path,
            mode=mode,
            engine=engine,
            include_candidates=True,
            offline_models=offline_models,
            require_models=require_models,
        )
        return result, output_path

    text = document.path.read_text(encoding="utf-8")
    result = humanize_text(
        text,
        mode=mode,
        engine=engine,
        include_candidates=True,
        offline_models=offline_models,
        require_models=require_models,
    )
    output_path = engine_dir / f"{document.id}.txt"
    output_path.write_text(result.text, encoding="utf-8")
    return result, output_path


def _row_from_payload(
    row: BenchmarkRow,
    *,
    payload: dict[str, Any],
    result: HumanizeResult,
    output_path: Path,
    report_path: Path,
    safety: dict[str, Any],
    status: str,
) -> BenchmarkRow:
    row.status = status
    row.output_path = str(output_path)
    row.report_path = str(report_path)
    row.accepted_changes = payload["summary"]["accepted_changes"]
    row.rejected_candidates = payload["summary"]["rejected_candidates"]
    row.skipped_sentences = payload["summary"]["skipped_sentences"]
    row.all_candidates = payload["summary"]["all_candidates"]
    row.changes_per_1000_words = payload["quality"]["changes_per_1000_words"]
    row.average_accepted_risk = payload["quality"]["average_accepted_risk"]
    row.operation_types = payload["quality"]["operation_types"]
    row.gate_rejections = payload["quality"]["gate_rejections"]
    row.model_status = result.model_status
    row.semantic_model = result.semantic_model
    row.fluency_model = result.fluency_model
    row.warnings = result.warnings
    row.safety = safety
    return row


def _write_summary_csv(rows: list[BenchmarkRow], path: Path) -> None:
    fieldnames = [
        "document_id",
        "document_type",
        "engine",
        "mode",
        "status",
        "accepted_changes",
        "rejected_candidates",
        "average_accepted_risk",
        "changes_per_1000_words",
        "safety_passed",
        "output_path",
        "report_path",
    ]
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow(
                {
                    "document_id": row.document_id,
                    "document_type": row.document_type,
                    "engine": row.engine,
                    "mode": row.mode,
                    "status": row.status,
                    "accepted_changes": row.accepted_changes,
                    "rejected_candidates": row.rejected_candidates,
                    "average_accepted_risk": row.average_accepted_risk,
                    "changes_per_1000_words": row.changes_per_1000_words,
                    "safety_passed": row.safety.get("passed"),
                    "output_path": row.output_path,
                    "report_path": row.report_path,
                }
            )


def _aggregate(rows: list[BenchmarkRow]) -> dict[str, Any]:
    return {
        "runs": len(rows),
        "ok": sum(1 for row in rows if row.status == "ok"),
        "failed_safety": sum(1 for row in rows if row.status == "failed_safety"),
        "model_unavailable": sum(1 for row in rows if row.status == "model_unavailable"),
        "accepted_changes": sum(row.accepted_changes for row in rows),
        "rejected_candidates": sum(row.rejected_candidates for row in rows),
        "average_risk": round(
            mean([row.average_accepted_risk for row in rows if row.accepted_changes]),
            4,
        )
        if any(row.accepted_changes for row in rows)
        else 0.0,
    }


def _rejected_sections(rows: list[BenchmarkRow]) -> list[str]:
    sections: list[str] = []
    for row in sorted(rows, key=lambda item: (item.document_id, item.engine)):
        if not row.report_path or not Path(row.report_path).exists():
            continue
        payload = json.loads(Path(row.report_path).read_text(encoding="utf-8"))
        rejected = payload.get("rejected", [])[:5]
        if not rejected:
            continue
        sections.append(f"### {row.document_id} / {row.engine}")
        for item in rejected:
            sections.append(
                f"- `{item.get('rule')}`: {item.get('reason')} "
                f"(sentence={item.get('sentence')})"
            )
    return sections


def _needs_review_sections(rows: list[BenchmarkRow]) -> list[str]:
    sections: list[str] = []
    for row in sorted(rows, key=lambda item: (item.document_id, item.engine)):
        if not row.report_path or not Path(row.report_path).exists():
            continue
        payload = json.loads(Path(row.report_path).read_text(encoding="utf-8"))
        risky = [
            item
            for item in payload.get("accepted", [])
            if (item.get("risk") or 0.0) >= 0.15
            or (item.get("semantic_similarity") is not None and item["semantic_similarity"] < 0.92)
            or (item.get("fluency_delta") is not None and item["fluency_delta"] < 0)
        ][:5]
        if not risky and row.safety.get("passed", True):
            continue
        sections.append(f"### {row.document_id} / {row.engine}")
        if not row.safety.get("passed", True):
            failed = [key for key, value in row.safety.items() if key != "passed" and not value]
            sections.append(f"- Safety failed: {', '.join(failed)}")
        for item in risky:
            sections.append(
                f"- `{item.get('rule')}` risk={item.get('risk')} "
                f"similarity={item.get('semantic_similarity')} fluency={item.get('fluency_delta')}"
            )
    return sections


def _recommended_rules(rows: list[BenchmarkRow]) -> list[str]:
    gate_counts: dict[str, int] = {}
    operation_counts: dict[str, int] = {}
    for row in rows:
        for name, count in row.gate_rejections.items():
            gate_counts[name] = gate_counts.get(name, 0) + count
        for name, count in row.operation_types.items():
            operation_counts[name] = operation_counts.get(name, 0) + count

    recommendations: list[str] = []
    if gate_counts.get("semantic_similarity", 0):
        recommendations.append("- Przejrzeć reguły odrzucane przez `semantic_similarity` i zawęzić ich kontekst.")
    if gate_counts.get("legal_anchor_retention", 0) or gate_counts.get("content_anchor_retention", 0):
        recommendations.append("- Dodać warianty reguł zachowujące kotwice treściowe i prawne.")
    if operation_counts.get("legal_ai_style_rewrite", 0) < 5:
        recommendations.append("- Rozbudować `legal_ai_style` o kolejne monotonne ramy AI-prawnicze.")
    if not recommendations:
        recommendations.append("- Brak dominującego wzorca; analizować ręcznie sekcję `Needs Review`.")
    return recommendations


def _docx_documents(paths: list[Path]) -> list[BenchmarkDocument]:
    documents: list[BenchmarkDocument] = []
    for path in paths:
        documents.append(
            BenchmarkDocument(
                id=path.stem,
                path=path,
                type="docx",
                focus=["manual docx"],
                source_kind="docx",
            )
        )
    return documents


def _read_document_text(document: BenchmarkDocument) -> str:
    if document.source_kind == "docx":
        return _read_docx_text(document.path)
    return document.path.read_text(encoding="utf-8")


def _read_docx_text(path: Path) -> str:
    from docx import Document  # type: ignore

    doc = Document(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def _numbers(text: str) -> list[str]:
    return re.findall(r"\d+(?:[,.]\d+)?", text)


def _protected_values(protected) -> set[str]:
    return set(protected.mapping.values())


def run_basic_tmp_benchmark(manifest: Path) -> Path:
    output = Path(tempfile.mkdtemp(prefix="humanize-pl-benchmark-"))
    rows = run_benchmark(
        load_manifest(manifest),
        output_dir=output,
        engines=[Engine.basic],
        mode=Mode.standard,
        offline_models=True,
        require_models=False,
        allow_fallback=True,
    )
    write_summary_artifacts(rows, output)
    return output


if __name__ == "__main__":
    app()
