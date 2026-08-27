"""End-to-end flow over a folder of .docx documents."""

from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Any

from humanize_pl.detect import detect_document
from humanize_pl.io.docx_io import docx_text
from .base import (
    FlowSettings,
    ItemOutcome,
    attach_pdf_report,
    layer_status,
    run_all_layers,
    summarise,
)


def docx_files(directory: Path) -> list[Path]:
    return sorted(
        (
            path
            for path in directory.iterdir()
            if path.is_file()
            and path.suffix.lower() == ".docx"
            and not path.name.startswith("~$")
        ),
        key=lambda path: (path.name.casefold(), path.name),
    )


def run_docx_flow(
    input_directory: Path,
    output_directory: Path,
    *,
    settings: FlowSettings,
    pdf: bool = True,
    on_item=None,
    on_layers=None,
) -> dict[str, Any]:
    """Diagnose, rewrite and gate every .docx in `input_directory`.

    One humanizer session is reused across documents so optional NLP models
    load once rather than per file.
    """
    files = docx_files(input_directory)
    if not files:
        raise FileNotFoundError(f"No .docx files in {input_directory}")

    output_directory.mkdir(parents=True, exist_ok=True)
    details_directory = output_directory / "details"
    details_directory.mkdir(parents=True, exist_ok=True)

    session = settings.session() if settings.rewrite else None
    layers = layer_status(session)
    if on_layers is not None:
        on_layers(layers)
    outcomes: list[ItemOutcome] = []

    for path in files:
        try:
            text = docx_text(path)
            outcome, verdict = run_all_layers(
                text, name=path.name, settings=settings, session=session
            )
            if settings.rewrite and outcome.text_out is not None:
                _write_docx(path, output_directory / f"{path.stem}_humanized.docx", outcome.text_out)
            _write_detail(details_directory / f"{path.stem}.json", text, outcome, verdict)
        except Exception as exc:  # one bad document must not stop the batch
            outcome = ItemOutcome(
                name=path.name, status="failed", error=f"{type(exc).__name__}: {exc}"
            )
        outcomes.append(outcome)
        if on_item is not None:
            on_item(outcome)

    summary = summarise(outcomes)
    payload = {
        "flow": "docx",
        "input_directory": str(input_directory),
        "output_directory": str(output_directory),
        "settings": {
            "mode": settings.mode.value,
            "engine": settings.engine.value,
            "rewrite": settings.rewrite,
            "require_anchor": settings.require_anchor,
        },
        "layers": layers,
        "summary": summary,
        "documents": [item.to_json() for item in outcomes],
    }
    if pdf:
        attach_pdf_report(payload, output_directory / "raport.pdf")
    (output_directory / "flow-report.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    _write_csv(output_directory / "summary.csv", outcomes)
    return payload


def _write_docx(source: Path, target: Path, text: str) -> None:
    """Rewrite paragraph text in place, preserving the original document."""
    from docx import Document  # type: ignore

    document = Document(str(source))
    lines = text.split("\n")
    index = 0
    for paragraph in document.paragraphs:
        if not paragraph.text.strip():
            continue
        if index < len(lines):
            if paragraph.text != lines[index]:
                paragraph.clear()
                paragraph.add_run(lines[index])
            index += 1
    document.save(str(target))


def _write_detail(path: Path, text: str, outcome: ItemOutcome, verdict) -> None:
    diagnosis = detect_document(text)
    path.write_text(
        json.dumps(
            {
                **outcome.to_json(),
                "findings": [
                    {
                        "family": finding.family,
                        "rule": finding.rule,
                        "evidence": finding.evidence,
                        "paragraph": finding.paragraph_index,
                        "sentence": finding.sentence_index,
                        "rewritable": finding.rewritable,
                    }
                    for finding in diagnosis.findings
                ],
                "metrics": diagnosis.metrics,
                "gate": verdict.to_json(),
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )


def _write_csv(path: Path, outcomes: list[ItemOutcome]) -> None:
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(
            [
                "dokument",
                "status",
                "slowa",
                "sygnal_przed",
                "sygnal_po",
                "delta",
                "do_przegladu",
                "znaleziska",
                "zmiany",
                "rodziny",
            ]
        )
        for item in outcomes:
            writer.writerow(
                [
                    item.name,
                    item.status,
                    item.words,
                    item.signal_before,
                    item.signal_after,
                    item.signal_delta,
                    "tak" if item.needs_review else "nie",
                    item.findings_before,
                    item.changes_applied,
                    "; ".join(item.families),
                ]
            )
