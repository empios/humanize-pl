from __future__ import annotations

from pathlib import Path

from humanize_pl.config import Engine, LegalReviewProfile, Mode
from humanize_pl.core import HumanizeResult, create_humanizer_session


def process_docx(
    input_path: str | Path,
    output_path: str | Path,
    *,
    mode: str | Mode = Mode.conservative,
    engine: str | Engine = Engine.basic,
    legal_review_profile: str | LegalReviewProfile = LegalReviewProfile.legal_ai_review,
    semantic_threshold: float | None = None,
    semantic_model: str | None = None,
    fluency_model: str | None = None,
    require_models: bool = False,
    offline_models: bool = False,
    include_candidates: bool = False,
    agreement_gate_enabled: bool = True,
    require_morfeusz: bool = False,
) -> tuple[HumanizeResult, dict[str, int]]:
    from docx import Document  # type: ignore

    input_path = Path(input_path)
    output_path = Path(output_path)
    doc = Document(str(input_path))

    processed = 0
    changed = 0
    empty = 0
    all_changes = []
    all_rejected = []
    all_skipped = []
    all_candidates = []
    session = create_humanizer_session(
        mode=mode,
        engine=engine,
        legal_review_profile=legal_review_profile,
        semantic_threshold=semantic_threshold,
        semantic_model=semantic_model,
        fluency_model=fluency_model,
        require_models=require_models,
        offline_models=offline_models,
        agreement_gate_enabled=agreement_gate_enabled,
        require_morfeusz=require_morfeusz,
    )
    warnings = list(session.warnings)
    engine_used = session.engine_used
    model_status: dict[str, str] = dict(session.model_status)
    semantic_model_used = (
        session.semantic_model_used if session.config.engine == Engine.hybrid else None
    )
    fluency_model_used = (
        session.fluency_model_used if session.config.engine == Engine.hybrid else None
    )

    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        original = paragraph.text
        if not original.strip():
            empty += 1
            continue
        processed += 1
        result = session.humanize(original, include_candidates=include_candidates)
        engine_used = result.engine_used
        model_status.update(result.model_status)
        semantic_model_used = result.semantic_model or semantic_model_used
        fluency_model_used = result.fluency_model or fluency_model_used
        for change in result.changes:
            change.paragraph_index = paragraph_index
        for rejection in result.rejected:
            rejection.paragraph_index = paragraph_index
        for skipped in result.skipped:
            skipped.paragraph_index = paragraph_index
        for candidate in result.all_candidates:
            candidate.paragraph_index = paragraph_index
        all_rejected.extend(result.rejected)
        all_skipped.extend(result.skipped)
        all_candidates.extend(result.all_candidates)
        if result.text != original:
            paragraph.clear()
            paragraph.add_run(result.text)
            changed += 1
            all_changes.extend(result.changes)

    doc.save(str(output_path))
    aggregate = HumanizeResult(
        text=str(output_path),
        changed=changed > 0,
        changes=all_changes,
        rejected=all_rejected,
        skipped=all_skipped,
        all_candidates=all_candidates,
        engine_used=engine_used,
        legal_review_profile=(
            legal_review_profile.value
            if isinstance(legal_review_profile, LegalReviewProfile)
            else str(legal_review_profile)
        ),
        model_status=model_status,
        semantic_model=semantic_model_used,
        fluency_model=fluency_model_used,
        warnings=warnings,
    )
    return aggregate, {"processed": processed, "changed": changed, "empty": empty}
