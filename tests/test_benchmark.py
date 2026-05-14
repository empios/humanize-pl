from __future__ import annotations

import json
from pathlib import Path

from typer.testing import CliRunner

import humanize_pl.benchmark as benchmark
from humanize_pl.benchmark import (
    BenchmarkDocument,
    BenchmarkRow,
    load_manifest,
    render_review_markdown,
    run_benchmark,
    safety_checks,
    write_summary_artifacts,
)
from humanize_pl.config import Engine, Mode
from humanize_pl.results import HumanizeResult


def test_manifest_loads_ai_legal_documents():
    documents = load_manifest(Path("docs_tests/ai_generated/manifest.json"))
    assert len(documents) == 8
    assert documents[0].id == "ai_legal_01_umowa_uslug"
    assert documents[-1].path.exists()


def test_manifest_rejects_missing_files(tmp_path):
    manifest = tmp_path / "manifest.json"
    manifest.write_text(
        json.dumps([{"id": "missing", "file": "missing.txt"}]),
        encoding="utf-8",
    )
    try:
        load_manifest(manifest)
    except FileNotFoundError as exc:
        assert "missing.txt" in str(exc)
    else:
        raise AssertionError("expected FileNotFoundError")


def test_safety_checks_detect_regressions():
    original = "W dniu 10.05.2026 r. strona zapłaciła 5000 zł."
    rewritten = "Surprisingly strona zapłaciła 6000 zł. __PROTECTED_0001__"
    checks = safety_checks(original, rewritten)
    assert not checks["passed"]
    assert not checks["no_english_markers"]
    assert not checks["no_placeholder_leak"]
    assert not checks["numbers_preserved"]


def test_review_markdown_contains_required_sections():
    rows = [
        BenchmarkRow(
            document_id="doc",
            document_type="umowa",
            engine="basic",
            mode="standard",
            status="ok",
            source_path="doc.txt",
            accepted_changes=2,
            rejected_candidates=1,
            average_accepted_risk=0.05,
            changes_per_1000_words=10.0,
            safety={"passed": True},
            gate_rejections={"semantic_similarity": 1},
        )
    ]
    markdown = render_review_markdown(rows)
    assert "## Summary" in markdown
    assert "## Per Document" in markdown
    assert "## Rejected Candidates" in markdown
    assert "## Needs Review" in markdown
    assert "## Recommended Next Rules" in markdown


def test_basic_benchmark_writes_all_artifacts(tmp_path):
    source_a = tmp_path / "a.txt"
    source_b = tmp_path / "b.txt"
    source_a.write_text("Podsumowując źródła prawa pracy tworzą system.", encoding="utf-8")
    source_b.write_text("Warto wskazać, że pracownik wykonuje pracę.", encoding="utf-8")
    manifest = tmp_path / "manifest.json"
    manifest.write_text(
        json.dumps(
            [
                {"id": "a", "file": "a.txt", "type": "essay"},
                {"id": "b", "file": "b.txt", "type": "letter"},
            ]
        ),
        encoding="utf-8",
    )
    output = tmp_path / "out"
    rows = run_benchmark(
        load_manifest(manifest),
        output_dir=output,
        engines=[Engine.basic],
        mode=Mode.standard,
        offline_models=True,
        require_models=False,
        allow_fallback=True,
    )
    write_summary_artifacts(rows, output)

    assert len(rows) == 2
    assert (output / "summary.json").exists()
    assert (output / "summary.csv").exists()
    assert (output / "review.md").exists()
    assert (output / "basic" / "a.txt").exists()
    assert (output / "basic" / "a.json").exists()


def test_long_docx_benchmark_records_timing_and_artifacts(monkeypatch, tmp_path):
    from docx import Document

    source = tmp_path / "long.docx"
    doc = Document()
    for index in range(60):
        doc.add_paragraph(
            f"Warto wskazać, że akapit {index} opisuje obowiązki strony umowy."
        )
    doc.save(source)

    def fake_process_docx(input_path, output_path, **kwargs):
        input_doc = Document(str(input_path))
        output_doc = Document()
        for paragraph in input_doc.paragraphs:
            output_doc.add_paragraph(
                paragraph.text.replace("Warto wskazać, że", "Wskazano, że")
            )
        output_doc.save(str(output_path))
        return (
            HumanizeResult(
                text=str(output_path),
                changed=True,
                engine_used="nlp",
                model_status={
                    "stanza": "ready",
                    "semantic": "not_requested",
                    "fluency": "not_requested",
                    "morfeusz": "ready",
                },
            ),
            {"processed": 60, "changed": 60, "empty": 0},
        )

    monkeypatch.setattr(benchmark, "process_docx", fake_process_docx)
    output = tmp_path / "out"
    rows = run_benchmark(
        [BenchmarkDocument(id="long_docx", path=source, type="docx", source_kind="docx")],
        output_dir=output,
        engines=[Engine.nlp],
        mode=Mode.standard,
        offline_models=True,
        require_models=True,
        allow_fallback=False,
    )
    write_summary_artifacts(rows, output)

    assert len(rows) == 1
    assert rows[0].status == "ok"
    assert rows[0].processing_seconds >= 0
    assert (output / "nlp" / "long_docx.docx").exists()
    assert (output / "nlp" / "long_docx.json").exists()
    review = (output / "review.md").read_text(encoding="utf-8")
    assert "Processing seconds" in review
    assert "Time(s)" in review


def test_fake_hybrid_records_model_metadata(monkeypatch, tmp_path):
    source = tmp_path / "a.txt"
    source.write_text("Pracownik wykonuje pracę.", encoding="utf-8")
    document = BenchmarkDocument(id="a", path=source, type="essay")

    def fake_humanize_text(*args, **kwargs):
        return HumanizeResult(
            text="Pracownik wykonuje pracę.",
            changed=False,
            engine_used="hybrid",
            model_status={"stanza": "ready", "semantic": "ready", "fluency": "ready"},
            semantic_model="fake-semantic",
            fluency_model="fake-fluency",
        )

    monkeypatch.setattr(benchmark, "humanize_text", fake_humanize_text)
    rows = run_benchmark(
        [document],
        output_dir=tmp_path / "out",
        engines=[Engine.hybrid],
        mode=Mode.standard,
        offline_models=True,
        require_models=True,
        allow_fallback=False,
    )
    assert rows[0].model_status["semantic"] == "ready"
    assert rows[0].semantic_model == "fake-semantic"


def test_benchmark_cli_help_exposes_options():
    runner = CliRunner()
    result = runner.invoke(benchmark.app, ["--help"])
    assert result.exit_code == 0
    assert "--engines" in result.stdout
    assert "--allow-fallback" in result.stdout
    assert "--fail-on-status" in result.stdout


def test_benchmark_cli_fail_on_status_exits_nonzero(monkeypatch, tmp_path):
    runner = CliRunner()

    def fake_load_manifest(path):
        return [BenchmarkDocument(id="doc", path=tmp_path / "doc.txt")]

    def fake_run_benchmark(*args, **kwargs):
        return [
            BenchmarkRow(
                document_id="doc",
                document_type="unknown",
                engine="basic",
                mode="standard",
                status="failed_safety",
                source_path="doc.txt",
                safety={"passed": False},
            )
        ]

    monkeypatch.setattr(benchmark, "load_manifest", fake_load_manifest)
    monkeypatch.setattr(benchmark, "run_benchmark", fake_run_benchmark)

    result = runner.invoke(
        benchmark.app,
        [
            "--manifest",
            str(tmp_path / "manifest.json"),
            "--output",
            str(tmp_path / "out"),
            "--engines",
            "basic",
            "--fail-on-status",
        ],
    )

    assert result.exit_code == 1
    assert "Statusy wymagające uwagi" in result.stdout
