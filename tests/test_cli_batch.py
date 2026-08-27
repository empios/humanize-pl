import json

from docx import Document
from typer.testing import CliRunner

from humanize_pl.cli import app


def _write_docx(path, *paragraphs: str) -> None:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


def test_cli_processes_docx_folder_and_writes_aggregate_report(tmp_path):
    input_directory = tmp_path / "input"
    output_directory = tmp_path / "output"
    report_path = tmp_path / "batch.json"
    input_directory.mkdir()
    _write_docx(input_directory / "b.docx", "Pracownik wykonuje pracę.")
    _write_docx(
        input_directory / "a.DOCX",
        "Podsumowując źródła prawa pracy tworzą system.",
    )
    (input_directory / "notes.txt").write_text("ignored", encoding="utf-8")
    (input_directory / "~$lock.docx").write_text("ignored", encoding="utf-8")

    result = CliRunner().invoke(
        app,
        [
            str(input_directory),
            "--output",
            str(output_directory),
            "--report",
            str(report_path),
            "--mode",
            "standard",
            "--no-agreement-gate",
        ],
    )

    assert result.exit_code == 0, result.stdout
    assert (output_directory / "a_humanized.docx").exists()
    assert (output_directory / "b_humanized.docx").exists()
    assert not (output_directory / "~$lock_humanized.docx").exists()

    output_document = Document(output_directory / "a_humanized.docx")
    assert output_document.paragraphs[0].text.startswith("Podsumowując,")

    payload = json.loads(report_path.read_text(encoding="utf-8"))
    assert payload["summary"] == {
        "documents": 2,
        "ok": 2,
        "failed": 0,
        "changed_documents": 1,
        "processed_paragraphs": 2,
        "changed_paragraphs": 1,
        "empty_paragraphs": 0,
        "accepted_changes": 1,
        "rejected_candidates": 0,
        "skipped_sentences": 1,
        "all_candidates": 0,
    }
    assert [document["input_path"] for document in payload["documents"]] == [
        str(input_directory / "a.DOCX"),
        str(input_directory / "b.docx"),
    ]
    assert all(document["report_summary"] for document in payload["documents"])
    assert payload["documents"][0]["report_summary"]["summary"]["accepted_changes"] == 1

    details_directory = tmp_path / "batch_details"
    assert (details_directory / "a.json").exists()
    assert (details_directory / "b.json").exists()
    assert payload["details_directory"] == str(details_directory)


def test_cli_folder_report_records_failures_and_continues(tmp_path):
    input_directory = tmp_path / "input"
    output_directory = tmp_path / "output"
    report_path = tmp_path / "batch.json"
    input_directory.mkdir()
    _write_docx(input_directory / "a_good.docx", "Pracownik wykonuje pracę.")
    (input_directory / "b_broken.docx").write_text("not a Word document", encoding="utf-8")

    result = CliRunner().invoke(
        app,
        [
            str(input_directory),
            "--output",
            str(output_directory),
            "--report",
            str(report_path),
            "--no-agreement-gate",
        ],
    )

    assert result.exit_code == 1
    assert (output_directory / "a_good_humanized.docx").exists()
    payload = json.loads(report_path.read_text(encoding="utf-8"))
    assert payload["summary"]["documents"] == 2
    assert payload["summary"]["ok"] == 1
    assert payload["summary"]["failed"] == 1
    assert payload["documents"][1]["status"] == "failed"
    assert "PackageNotFoundError" in payload["documents"][1]["error"]


def test_cli_rejects_folder_without_docx_files(tmp_path):
    input_directory = tmp_path / "input"
    input_directory.mkdir()
    (input_directory / "notes.txt").write_text("text", encoding="utf-8")

    result = CliRunner().invoke(app, [str(input_directory)])

    assert result.exit_code == 2
    assert "does not contain any .docx files" in result.stderr
